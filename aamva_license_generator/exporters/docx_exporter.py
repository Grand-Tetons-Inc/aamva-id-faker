"""
DOCX Export using python-docx

Generates Word documents with Avery 28371 business card layout.
Extracted and improved from generate_licenses.py with:
- Better error handling
- Progress tracking
- Resource management
- Proper cleanup
"""

from pathlib import Path
from typing import List, Dict, Any, Tuple

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError as e:
    raise ImportError(
        "python-docx is required for DOCX export. Install with: pip install python-docx"
    ) from e

from .base import (
    BaseExporter, ExportFormat, ExportOptions, ExportResult,
    ValidationError, RenderError
)
from ..storage import (
    FileSystemValidator, SafeFileOperations, StorageError,
    TemporaryFileManager
)


class DOCXExporter(BaseExporter):
    """
    Export license data to Word document using Avery 28371 layout

    Layout specifications:
    - Page: 8.5" x 11" (US Letter)
    - Cards per page: 10 (2 columns x 5 rows)
    - Card size: 3.5" x 2" (business card)
    - Margins: 0.75" left/right, 0.5" top/bottom
    """

    # Layout constants
    PAGE_WIDTH = 8.5
    PAGE_HEIGHT = 11
    CARD_WIDTH = 3.5
    CARD_HEIGHT = 2.0
    CARDS_PER_PAGE = 10

    @property
    def format(self) -> ExportFormat:
        return ExportFormat.DOCX

    @property
    def file_extension(self) -> str:
        return "docx"

    def validate_data(self, data: Any) -> None:
        """
        Validate license data for DOCX export

        Args:
            data: List of (barcode_path, license_data) tuples

        Raises:
            ValidationError: If data is invalid
        """
        if not isinstance(data, list):
            raise ValidationError("Data must be a list of license records")

        if len(data) == 0:
            raise ValidationError("No license data to export")

        for index, item in enumerate(data):
            if not isinstance(item, (tuple, list)) or len(item) != 2:
                raise ValidationError(
                    f"Item {index}: Must be (barcode_path, license_data) tuple"
                )

            barcode_path, license_data = item

            # Validate barcode path
            if not isinstance(barcode_path, (str, Path)):
                raise ValidationError(
                    f"Item {index}: Barcode path must be string or Path"
                )

            # Note: We don't require barcode to exist since we'll generate card images

            # Validate license data
            if not isinstance(license_data, list) or len(license_data) < 1:
                raise ValidationError(
                    f"Item {index}: License data must be list with at least one subfile"
                )

    def _export_impl(self, data: List[Tuple[Path, List[Dict[str, Any]]]]) -> ExportResult:
        """
        Export licenses to DOCX

        Args:
            data: List of (barcode_path, license_data) tuples

        Returns:
            ExportResult with operation details
        """
        result = ExportResult(success=True)
        output_path = Path(self.options.output_path)

        try:
            # Ensure output directory exists
            if not output_path.parent.exists():
                raise RenderError(f"Output directory does not exist: {output_path.parent}")

            # Check writable
            if not FileSystemValidator.check_writable(output_path.parent):
                raise RenderError(f"Cannot write to directory: {output_path.parent}")

            # Estimate size and check disk space
            estimated_size = self._estimate_docx_size(len(data))
            FileSystemValidator.ensure_space(output_path, estimated_size)

            # Generate DOCX
            with SafeFileOperations.atomic_write(output_path, mode='wb') as f:
                self._generate_docx(f.name, data)

            result.output_path = output_path
            result.items_processed = len(data)

        except (StorageError, RenderError) as e:
            result.success = False
            result.errors.append(str(e))
        except Exception as e:
            result.success = False
            result.errors.append(f"Unexpected error: {e}")

        return result

    def _estimate_docx_size(self, num_licenses: int) -> int:
        """
        Estimate DOCX file size

        Args:
            num_licenses: Number of licenses

        Returns:
            Estimated size in bytes
        """
        # Rough estimate: 100KB base + 50KB per license (includes images)
        return 100 * 1024 + (num_licenses * 50 * 1024)

    def _generate_docx(self, filepath: str, data: List[Tuple[Path, List[Dict[str, Any]]]]):
        """
        Generate the actual DOCX file

        Args:
            filepath: Output file path
            data: License data
        """
        total = len(data)
        doc = None
        card_images = []  # Track temporary card images for cleanup

        try:
            # Create document
            self._update_progress(0, total, "initializing", "Creating document...")
            doc = Document()

            # Set page size and margins
            section = doc.sections[-1]
            section.page_width = Inches(self.PAGE_WIDTH)
            section.page_height = Inches(self.PAGE_HEIGHT)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)

            # Create card table
            table = self._add_card_table(doc)
            cards_in_page = 0

            # Process each card
            for idx, (img_path, license_data) in enumerate(data):
                self._update_progress(
                    idx,
                    total,
                    "rendering",
                    f"Processing card {idx + 1} of {total}..."
                )

                # Generate card image
                try:
                    card_img_path = self._generate_card_image(img_path, license_data, idx)
                    card_images.append(card_img_path)
                except Exception as e:
                    result_error = f"Failed to generate card image {idx}: {e}"
                    print(f"Warning: {result_error}")
                    card_img_path = None

                # Calculate cell position
                row_idx = (cards_in_page) // 2
                col_idx = cards_in_page % 2

                # Add card to table
                try:
                    cell = table.cell(row_idx, col_idx)
                    if card_img_path:
                        self._add_card_to_cell(cell, card_img_path)
                    else:
                        self._add_error_to_cell(cell, f"Card {idx}: Image generation failed")
                except Exception as e:
                    print(f"Warning: Failed to add card {idx} to table: {e}")

                cards_in_page += 1

                # Add new page after 10 cards (except for last card)
                if cards_in_page >= self.CARDS_PER_PAGE and idx < total - 1:
                    table = self._add_card_table(doc)
                    cards_in_page = 0

            # Save document
            self._update_progress(total, total, "saving", "Saving document...")
            doc.save(filepath)

        except Exception as e:
            raise RenderError(f"Failed to generate DOCX: {e}") from e

        finally:
            # Clean up temporary card images
            for img_path in card_images:
                try:
                    if img_path and Path(img_path).exists():
                        Path(img_path).unlink()
                except Exception as e:
                    print(f"Warning: Could not clean up temporary image {img_path}: {e}")

    def _add_card_table(self, doc: Document):
        """
        Add a card table to the document

        Args:
            doc: Document object

        Returns:
            Table object
        """
        table = doc.add_table(rows=5, cols=2)
        table.autofit = False

        # Set cell dimensions
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()

                # Set cell width
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(self.CARD_WIDTH * 1440)))  # 1440 twips per inch
                tcW.set(qn('w:type'), 'dxa')
                tc_pr.append(tcW)

                # Remove cell margins
                tcMar = OxmlElement('w:tcMar')
                for margin_type in ['top', 'left', 'bottom', 'right']:
                    node = OxmlElement(f'w:{margin_type}')
                    node.set(qn('w:w'), '0')
                    node.set(qn('w:type'), 'dxa')
                    tcMar.append(node)
                tc_pr.append(tcMar)

                cell.width = Inches(self.CARD_WIDTH)

            # Set row height
            tr_pr = row._tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(int(self.CARD_HEIGHT * 1440)))
            trHeight.set(qn('w:hRule'), 'exact')
            tr_pr.append(trHeight)
            row.height = Inches(self.CARD_HEIGHT)

        return table

    def _generate_card_image(self, barcode_path: Path,
                            license_data: List[Dict[str, Any]],
                            index: int) -> Path:
        """
        Generate a card image from barcode and license data

        This is a simplified version. In production, you'd use the
        generate_individual_card_image function from generate_licenses.py

        Args:
            barcode_path: Path to barcode image
            license_data: License data
            index: Card index

        Returns:
            Path to generated card image
        """
        from PIL import Image as PILImage, ImageDraw, ImageFont

        # Import the actual implementation from generate_licenses
        # For now, create a simple placeholder
        try:
            # Create card image
            dpi = self.options.dpi
            card_width = int(self.CARD_WIDTH * dpi)
            card_height = int(self.CARD_HEIGHT * dpi)

            card = PILImage.new("RGB", (card_width, card_height), "white")
            draw = ImageDraw.Draw(card)

            # Load barcode if exists
            if Path(barcode_path).exists():
                barcode_img = PILImage.open(barcode_path).convert("L")
                barcode_width = int(card_width * 0.55)
                barcode_height = int(barcode_width * 0.25)
                barcode_img = barcode_img.resize(
                    (barcode_width, barcode_height),
                    PILImage.Resampling.LANCZOS
                )

                barcode_x = int(card_width * 0.03)
                barcode_y = int(card_height * 0.05)
                card.paste(barcode_img, (barcode_x, barcode_y))

                # Add text below barcode
                text_y = barcode_y + barcode_height + int(card_height * 0.02)
            else:
                text_y = int(card_height * 0.1)

            # Add license data text
            try:
                font_size = int(40 * (dpi / 300))  # Scale with DPI
                font = ImageFont.truetype("LiberationMono-Bold.ttf", font_size)
            except Exception:
                font = ImageFont.load_default()

            dl_data = license_data[0]
            text_x = int(card_width * 0.03)

            # Format text
            lines = [
                f"{dl_data.get('DAC', '')} {dl_data.get('DAD', '')} {dl_data.get('DCS', '')}",
                f"DOB: {dl_data.get('DBB', '')} | EXP: {dl_data.get('DBA', '')}",
                f"DL#: {dl_data.get('DAQ', '')}",
                f"Class: {dl_data.get('DCA', '')} | {dl_data.get('DAI', '')}, {dl_data.get('DAJ', '')}",
            ]

            text = "\n".join(lines)
            draw.text((text_x, text_y), text, fill="black", font=font, spacing=10)

            # Save to temporary file
            with TemporaryFileManager.temporary_file(
                suffix=f"_card_{index}.png",
                cleanup=False  # We'll clean up manually
            ) as (f, temp_path):
                card.save(temp_path, dpi=(dpi, dpi))
                return temp_path

        except Exception as e:
            raise RenderError(f"Failed to generate card image: {e}") from e

    def _add_card_to_cell(self, cell, card_img_path: Path):
        """
        Add card image to table cell

        Args:
            cell: Table cell
            card_img_path: Path to card image
        """
        # Clear existing content
        cell.text = ""
        paragraph = cell.paragraphs[0]

        # Add image
        run = paragraph.add_run()
        run.add_picture(str(card_img_path), width=Inches(3.4))

        # Center the image
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Remove spacing
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing = 1.0

    def _add_error_to_cell(self, cell, error_message: str):
        """
        Add error message to table cell

        Args:
            cell: Table cell
            error_message: Error message to display
        """
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(f"ERROR\n\n{error_message}")
        run.font.size = Pt(8)
        run.font.color.rgb = (200, 0, 0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


class DOCXExportOptions(ExportOptions):
    """Extended options specific to DOCX export"""

    def __init__(self, output_path: str, **kwargs):
        super().__init__(output_path, **kwargs)
        self.page_width: float = 8.5
        self.page_height: float = 11.0
        self.left_margin: float = 0.75
        self.right_margin: float = 0.75
        self.top_margin: float = 0.5
        self.bottom_margin: float = 0.5
        self.card_image_width: float = 3.4  # Slightly less than card width for margins
