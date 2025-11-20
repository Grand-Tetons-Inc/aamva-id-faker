"""
Export Service - License Data Export Orchestration

Handles exporting license data to various formats:
- PDF417 barcodes (BMP/PNG)
- AAMVA text format
- PDF documents (Avery templates)
- DOCX documents
- JSON/CSV data exports
"""

import os
import json
import csv
from typing import Dict, List, Optional, Callable, Any, Tuple
from pathlib import Path
import logging

try:
    import pdf417
    from PIL import Image as PILImage, ImageDraw, ImageFont
except ImportError:
    pdf417 = None
    PILImage = None

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import inch
except ImportError:
    letter = None
    canvas = None

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    Document = None

# Configure logging
logger = logging.getLogger(__name__)


class ExportError(Exception):
    """Raised when export operation fails"""
    pass


class ExportResult:
    """
    Result of an export operation.

    Attributes:
        success: Whether export succeeded
        output_paths: List of generated file paths
        errors: List of error messages
        warnings: List of warning messages
    """

    def __init__(self):
        self.success = True
        self.output_paths: List[str] = []
        self.errors: List[str] = []
        self.warnings: List[str] = []

    def add_path(self, path: str):
        """Add an output file path."""
        self.output_paths.append(path)

    def add_error(self, message: str):
        """Add an error message."""
        self.success = False
        self.errors.append(message)

    def add_warning(self, message: str):
        """Add a warning message."""
        self.warnings.append(message)

    def __str__(self) -> str:
        """String representation."""
        if self.success:
            return f"✓ Export successful: {len(self.output_paths)} files"
        else:
            return f"✗ Export failed: {len(self.errors)} errors"

    def __bool__(self) -> bool:
        """Boolean representation."""
        return self.success


class ExportService:
    """
    Service for exporting license data to various formats.

    Supports:
    - Barcode generation (BMP, PNG)
    - AAMVA text format
    - PDF documents (Avery templates)
    - DOCX documents
    - JSON/CSV data files
    """

    def __init__(self, output_dir: str = "output"):
        """
        Initialize the ExportService.

        Args:
            output_dir: Base directory for exports
        """
        self.output_dir = output_dir
        self.barcode_dir = os.path.join(output_dir, "barcodes")
        self.data_dir = os.path.join(output_dir, "data")
        self.cards_dir = os.path.join(output_dir, "cards")
        self._ensure_directories()

    def _ensure_directories(self):
        """Create output directories if they don't exist."""
        try:
            os.makedirs(self.barcode_dir, exist_ok=True)
            os.makedirs(self.data_dir, exist_ok=True)
            os.makedirs(self.cards_dir, exist_ok=True)
        except Exception as e:
            raise ExportError(f"Failed to create directories: {e}") from e

    def get_iin_by_state(self, state_abbr: str) -> str:
        """
        Get IIN code for a state.

        Args:
            state_abbr: Two-letter state abbreviation

        Returns:
            IIN code (defaults to Arizona if not found)
        """
        # Minimal IIN mapping (full mapping should come from config/service)
        iin_map = {
            "CA": "636014", "TX": "636015", "FL": "636010", "NY": "636001",
            "PA": "636025", "IL": "636035", "OH": "636023", "GA": "636055",
            "NC": "636004", "MI": "636032", "NJ": "636036", "VA": "636000",
            "WA": "636045", "AZ": "636026", "MA": "636002", "TN": "636053",
        }
        return iin_map.get(state_abbr.upper(), "636026")  # Default to AZ

    def format_barcode_data(self, license_data: List[Dict[str, str]]) -> str:
        """
        Format license data as AAMVA barcode string.

        Args:
            license_data: List containing [DL_data, State_data]

        Returns:
            AAMVA formatted barcode string

        Raises:
            ExportError: If formatting fails
        """
        try:
            # AAMVA compliance markers
            compliance = "@\n\x1E\r"
            file_type = "ANSI "
            version = "10"  # AAMVA 2020
            jurisdiction_version = "00"
            number_of_entries = f"{len(license_data):02d}"

            # Get DL and state data
            dl_data = license_data[0]
            state_data = license_data[1] if len(license_data) > 1 else {}

            # Get IIN
            iin = self.get_iin_by_state(dl_data.get("DAJ", "AZ"))

            # Build header base
            header_base = (
                compliance +
                file_type +
                iin +
                version +
                jurisdiction_version +
                number_of_entries
            )

            # Build DL subfile
            daq = dl_data.get("DAQ", "")
            dl_subfile_type = dl_data.get("subfile_type", "DL")
            dl_fields = {k: v for k, v in dl_data.items()
                        if k not in ["DAQ", "subfile_type"]}
            dl_subfile_data = (
                dl_subfile_type +
                f"DAQ{daq}\n" +
                "".join(f"{k}{v}\n" for k, v in dl_fields.items()) +
                "\r"
            )

            # Calculate offsets and lengths
            subfiles_designators_len = 10 * int(number_of_entries)
            dl_subfile_offset = len(header_base.encode("ascii")) + subfiles_designators_len
            dl_subfile_length = len(dl_subfile_data.encode("ascii"))
            dl_subfile_designator = (
                dl_subfile_type +
                f"{dl_subfile_offset:04d}" +
                f"{dl_subfile_length:04d}"
            )

            # Build state subfile if present
            if state_data:
                state_fields = {k: v for k, v in state_data.items()
                              if k != "subfile_type"}
                state_subfile_type = state_data.get("subfile_type", "ZX")
                state_subfile_data = (
                    state_subfile_type +
                    "".join(f"{k}{v}\n" for k, v in state_fields.items()) +
                    "\r"
                )
                state_subfile_length = len(state_subfile_data.encode("ascii"))
                state_subfile_offset = dl_subfile_offset + dl_subfile_length
                state_subfile_designator = (
                    state_subfile_type +
                    f"{state_subfile_offset:04d}" +
                    f"{state_subfile_length:04d}"
                )
                header = header_base + dl_subfile_designator + state_subfile_designator
                return header + dl_subfile_data + state_subfile_data
            else:
                header = header_base + dl_subfile_designator
                return header + dl_subfile_data

        except Exception as e:
            raise ExportError(f"Failed to format barcode data: {e}") from e

    def export_barcode(
        self,
        license_data: List[Dict[str, str]],
        index: int,
        format: str = "bmp"
    ) -> Tuple[str, str]:
        """
        Export license as barcode image.

        Args:
            license_data: License data array
            index: File index number
            format: Image format ('bmp' or 'png')

        Returns:
            Tuple of (image_path, data_path)

        Raises:
            ExportError: If export fails
        """
        if pdf417 is None or PILImage is None:
            raise ExportError("pdf417 and Pillow libraries required for barcode export")

        try:
            # Format barcode data
            raw = self.format_barcode_data(license_data)

            # Encode as PDF417
            codes = pdf417.encode(raw, columns=13, security_level=5)
            image = pdf417.render_image(codes)

            # Save barcode image
            img_path = os.path.join(self.barcode_dir, f"license_{index}.{format}")
            if format.lower() == "png":
                image = image.convert("RGB")
            image.save(img_path)

            # Save raw data
            txt_path = os.path.join(self.data_dir, f"license_{index}.txt")
            with open(txt_path, "w") as f:
                f.write(raw)

            logger.info(f"Exported barcode {index}: {img_path}")
            return img_path, txt_path

        except Exception as e:
            raise ExportError(f"Failed to export barcode {index}: {e}") from e

    def export_json(
        self,
        licenses: List[List[Dict[str, str]]],
        filename: str = "licenses.json"
    ) -> str:
        """
        Export licenses as JSON file.

        Args:
            licenses: List of license data arrays
            filename: Output filename

        Returns:
            Output file path

        Raises:
            ExportError: If export fails
        """
        try:
            output_path = os.path.join(self.output_dir, filename)

            with open(output_path, 'w') as f:
                json.dump(licenses, f, indent=2)

            logger.info(f"Exported JSON: {output_path}")
            return output_path

        except Exception as e:
            raise ExportError(f"Failed to export JSON: {e}") from e

    def export_csv(
        self,
        licenses: List[List[Dict[str, str]]],
        filename: str = "licenses.csv"
    ) -> str:
        """
        Export licenses as CSV file (flattened DL data).

        Args:
            licenses: List of license data arrays
            filename: Output filename

        Returns:
            Output file path

        Raises:
            ExportError: If export fails
        """
        try:
            output_path = os.path.join(self.output_dir, filename)

            if not licenses:
                raise ExportError("No licenses to export")

            # Get all possible fields from first license
            dl_data = licenses[0][0]
            fieldnames = [k for k in dl_data.keys() if k != "subfile_type"]

            with open(output_path, 'w', newline='') as f:
                writer = csv.DictWriter(f, fieldnames=fieldnames)
                writer.writeheader()

                for license_data in licenses:
                    dl_data = license_data[0]
                    row = {k: v for k, v in dl_data.items() if k != "subfile_type"}
                    writer.writerow(row)

            logger.info(f"Exported CSV: {output_path}")
            return output_path

        except Exception as e:
            raise ExportError(f"Failed to export CSV: {e}") from e

    def generate_card_image(
        self,
        license_data: List[Dict[str, str]],
        barcode_path: str,
        width_inches: float = 3.5,
        dpi: int = 300
    ) -> str:
        """
        Generate an individual card image.

        Args:
            license_data: License data array
            barcode_path: Path to barcode image
            width_inches: Card width in inches
            dpi: Dots per inch

        Returns:
            Path to generated card image

        Raises:
            ExportError: If generation fails
        """
        if PILImage is None:
            raise ExportError("Pillow library required for card image generation")

        try:
            card_width = int(width_inches * dpi)
            card_height = int(2.0 * dpi)

            card = PILImage.new("RGB", (card_width, card_height), "white")
            draw = ImageDraw.Draw(card)

            # Font sizes
            base_font_size = 50
            small_font_size = 40

            # Load font
            try:
                font = ImageFont.truetype("LiberationMono-Bold.ttf", base_font_size)
                small_font = ImageFont.truetype("LiberationMono-Bold.ttf", small_font_size)
            except:
                font = ImageFont.load_default()
                small_font = font

            # Add barcode
            barcode_img = PILImage.open(barcode_path).convert("L")
            barcode_width = int(card_width * 0.55)
            barcode_height = int(barcode_width * 0.25)

            barcode_img = barcode_img.resize(
                (barcode_width, barcode_height),
                PILImage.Resampling.LANCZOS
            )

            barcode_x = int(card_width * 0.03)
            barcode_y = int(card_height * 0.05)
            card.paste(barcode_img, (barcode_x, barcode_y))

            # Add text
            text_x = barcode_x
            text_y = barcode_y + barcode_height + int(card_height * 0.02)

            dl_data = license_data[0]
            state_data = license_data[1] if len(license_data) > 1 else {}

            state_line = "|".join([
                f"{key} {value}"
                for key, value in state_data.items()
                if key != "subfile_type"
            ])

            lines = (
                f"{dl_data.get('DAC', '')} {dl_data.get('DAD', '')} {dl_data.get('DCS', '')}\n" +
                f"DOB: {dl_data.get('DBB', '')} | EXP: {dl_data.get('DBA', '')}\n" +
                f"DL#: {dl_data.get('DAQ', '')}\n" +
                f"Class: {dl_data.get('DCA', '')} | {dl_data.get('DAI', '')}, {dl_data.get('DAJ', '')}\n" +
                f"{'M' if dl_data.get('DBC') == '1' else 'F'} | {dl_data.get('DAY', '')} | " +
                f"{dl_data.get('DAZ', '')} | {dl_data.get('DAU', '')}\" | {dl_data.get('DAW', '')}lbs\n" +
                f"Organ Donor: {dl_data.get('DDK', '')} | Veteran: {dl_data.get('DDL', '')}\n" +
                f"{state_line}"
            )

            draw.text((text_x, text_y), lines, fill="black", font=small_font, spacing=10)

            # Save card image
            card_img_path = barcode_path.replace('.bmp', '_card.png').replace('.png', '_card.png')
            if not card_img_path.endswith('_card.png'):
                card_img_path = card_img_path + '_card.png'

            card.save(card_img_path, dpi=(dpi, dpi))
            logger.info(f"Generated card image: {card_img_path}")
            return card_img_path

        except Exception as e:
            raise ExportError(f"Failed to generate card image: {e}") from e

    def export_pdf(
        self,
        records: List[Tuple[str, List[Dict[str, str]]]],
        filename: str = "licenses_avery_28371.pdf",
        progress_callback: Optional[Callable[[int, int], None]] = None
    ) -> str:
        """
        Export licenses as PDF with Avery 28371 layout.

        Args:
            records: List of (barcode_path, license_data) tuples
            filename: Output filename
            progress_callback: Optional progress callback(current, total)

        Returns:
            Output file path

        Raises:
            ExportError: If export fails
        """
        if canvas is None:
            raise ExportError("reportlab library required for PDF export")

        try:
            output_path = os.path.join(self.output_dir, filename)

            # Avery 28371 specifications
            page_width, page_height = letter
            card_width = 3.5 * inch
            card_height = 2 * inch
            left_margin = 0.75 * inch
            top_margin = 0.5 * inch
            horizontal_spacing = 0.25 * inch
            vertical_spacing = 0 * inch

            # Create PDF
            c = canvas.Canvas(output_path, pagesize=letter)

            # Process cards (10 per page)
            total = len(records)
            for page_num in range(0, total, 10):
                page_cards = records[page_num:page_num + 10]

                for card_index, (img_path, data) in enumerate(page_cards):
                    # Calculate position (2 columns, 5 rows)
                    row = card_index // 2
                    col = card_index % 2

                    x = left_margin + col * (card_width + horizontal_spacing)
                    y = page_height - top_margin - (row + 1) * (card_height + vertical_spacing)

                    # Draw border
                    c.setStrokeColorRGB(0.8, 0.8, 0.8)
                    c.setLineWidth(0.5)
                    c.rect(x, y, card_width, card_height, stroke=1, fill=0)

                    # Add barcode
                    barcode_width = 1.8 * inch
                    barcode_height = 0.6 * inch
                    barcode_x = x + 0.1 * inch
                    barcode_y = y + card_height - barcode_height - 0.1 * inch

                    try:
                        c.drawImage(img_path, barcode_x, barcode_y,
                                  width=barcode_width, height=barcode_height)
                    except:
                        logger.warning(f"Could not add barcode image: {img_path}")

                    # Add text
                    text_x = x + 0.1 * inch
                    text_y = y + card_height - barcode_height - 0.25 * inch

                    c.setFont("Helvetica", 8)
                    c.setFillColorRGB(0, 0, 0)

                    dl_data = data[0]
                    state_data = data[1] if len(data) > 1 else {}

                    state_line = "|".join([
                        f"{key} {value}"
                        for key, value in state_data.items()
                        if key != "subfile_type"
                    ])

                    lines = [
                        f"{dl_data.get('DAC', '')} {dl_data.get('DAD', '')} {dl_data.get('DCS', '')}",
                        f"DOB: {dl_data.get('DBB', '')} | EXP: {dl_data.get('DBA', '')}",
                        f"DL#: {dl_data.get('DAQ', '')}",
                        f"Class: {dl_data.get('DCA', '')} | {dl_data.get('DAI', '')}, {dl_data.get('DAJ', '')}",
                        f"{'M' if dl_data.get('DBC') == '1' else 'F'} | {dl_data.get('DAY', '')} | " +
                        f"{dl_data.get('DAZ', '')} | {dl_data.get('DAU', '')}\" | {dl_data.get('DAW', '')}lbs",
                        f"Organ Donor: {dl_data.get('DDK', '')} | Veteran: {dl_data.get('DDL', '')}",
                        f"{state_line}"
                    ]

                    line_height = 0.15 * inch
                    for i, line in enumerate(lines):
                        c.drawString(text_x, text_y - i * line_height, line)

                    if progress_callback:
                        progress_callback(page_num + card_index + 1, total)

                # Start new page if more cards
                if page_num + 10 < total:
                    c.showPage()

            c.save()
            logger.info(f"Exported PDF: {output_path}")
            return output_path

        except Exception as e:
            raise ExportError(f"Failed to export PDF: {e}") from e

    def export_docx(
        self,
        records: List[Tuple[str, List[Dict[str, str]]]],
        filename: str = "licenses.docx",
        progress_callback: Optional[Callable[[int, int], None]] = None
    ) -> str:
        """
        Export licenses as DOCX with table layout.

        Args:
            records: List of (barcode_path, license_data) tuples
            filename: Output filename
            progress_callback: Optional progress callback(current, total)

        Returns:
            Output file path

        Raises:
            ExportError: If export fails
        """
        if Document is None:
            raise ExportError("python-docx library required for DOCX export")

        try:
            output_path = os.path.join(self.output_dir, filename)

            doc = Document()
            section = doc.sections[-1]
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)

            card_width = 3.5
            card_height = 2.0
            max_cards_per_page = 10

            def add_card_table():
                table = doc.add_table(rows=5, cols=2)
                table.autofit = False

                for row in table.rows:
                    for cell in row.cells:
                        tc_pr = cell._tc.get_or_add_tcPr()

                        # Set cell width
                        tcW = OxmlElement('w:tcW')
                        tcW.set(qn('w:w'), str(int(card_width * 1440)))
                        tcW.set(qn('w:type'), 'dxa')
                        tc_pr.append(tcW)

                        # Remove margins
                        tcMar = OxmlElement('w:tcMar')
                        for margin_type in ['top', 'left', 'bottom', 'right']:
                            node = OxmlElement(f'w:{margin_type}')
                            node.set(qn('w:w'), '0')
                            node.set(qn('w:type'), 'dxa')
                            tcMar.append(node)
                        tc_pr.append(tcMar)

                        cell.width = Inches(card_width)

                    # Set row height
                    tr_pr = row._tr.get_or_add_trPr()
                    trHeight = OxmlElement('w:trHeight')
                    trHeight.set(qn('w:val'), str(int(card_height * 1440)))
                    trHeight.set(qn('w:hRule'), 'exact')
                    tr_pr.append(trHeight)
                    row.height = Inches(card_height)

                return table

            table = add_card_table()
            total = len(records)

            for idx, (img_path, data) in enumerate(records):
                # Calculate cell position
                row_idx = (idx % 10) // 2
                col_idx = idx % 2
                cell = table.cell(row_idx, col_idx)

                cell.text = ""
                paragraph = cell.paragraphs[0]

                # Generate card image
                card_img_path = self.generate_card_image(data, img_path, width_inches=3.4, dpi=300)

                # Add image
                run = paragraph.add_run()
                run.add_picture(card_img_path, width=Inches(3.4))

                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_before = Pt(0)
                paragraph_format.space_after = Pt(0)
                paragraph_format.line_spacing = 1.0

                if progress_callback:
                    progress_callback(idx + 1, total)

                # New page after 10 cards
                if (idx + 1) % max_cards_per_page == 0 and idx != total - 1:
                    table = add_card_table()

            doc.save(output_path)
            logger.info(f"Exported DOCX: {output_path}")
            return output_path

        except Exception as e:
            raise ExportError(f"Failed to export DOCX: {e}") from e
