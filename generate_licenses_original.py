#!/bin/sh
"true" '''\'
exec "$(dirname "$(readlink -f "$0")")"/.venv/bin/python "$0" "$@"
'''

# -----------------------------------------------------------------------------
# MIT License
# 
# Copyright (c) 2025 James W Rogers, Jr. / Grand Tetons Inc.
# Written for SMLA Inc.
# 
# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the "Software"), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:
# 
# The above copyright notice and this permission notice shall be included in all
# copies or substantial portions of the Software.
# 
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
# IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
# FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
# AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
# LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
# OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
# SOFTWARE.
# -----------------------------------------------------------------------------

import os
import argparse
import random
import string
from faker import Faker
from datetime import datetime, timedelta
import pdf417
from PIL import Image as PILImage, ImageDraw, ImageFont
from odf.opendocument import OpenDocumentText
from odf.style import Style, TextProperties, ParagraphProperties, SectionProperties, Columns, Column
from odf.text import P, Section, Span
from odf.table import Table, TableRow, TableCell
from odf.draw import Frame, Image as ODFImage
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from reportlab.platypus import Image as RLImage
from reportlab.lib.colors import black, lightgrey
import tempfile

fake = Faker()

# === CONFIG ===
# Required packages:
# pip install faker pdf417 pillow odfpy reportlab python-docx
OUTPUT_DIR = "output"
BARCODE_DIR = os.path.join(OUTPUT_DIR, "barcodes")
DATA_DIR = os.path.join(OUTPUT_DIR, "data")
CARDS_DIR = os.path.join(OUTPUT_DIR, "cards")
ODT_FILE = os.path.join(OUTPUT_DIR, "cards.odt")
DOCX_FILE = os.path.join(OUTPUT_DIR, "cards.docx")
#IINs: IINs: Issuer ID Numbers: 
# https://gist.github.com/ix4/1351ffda0137c5de015b7e710e486902
# https://www.aamva.org/identity/issuer-identification-numbers-(iin)
IIN_JURISDICTIONS = {
    "604426": {"jurisdiction": "Prince Edward Island", "abbr": "PE", "country": "Canada"},
    "604427": {"jurisdiction": "American Samoa", "abbr": "AS", "country": "USA"},
    "604428": {"jurisdiction": "Quebec", "abbr": "QC", "country": "Canada"},
    "604429": {"jurisdiction": "Yukon", "abbr": "YT", "country": "Canada"},
    "604430": {"jurisdiction": "Northern Marianna Islands", "abbr": "MP", "country": "USA"},
    "604431": {"jurisdiction": "Puerto Rico", "abbr": "PR", "country": "USA"},
    "604432": {"jurisdiction": "Alberta", "abbr": "AB", "country": "Canada"},
    "604433": {"jurisdiction": "Nunavut", "abbr": "NU", "country": "Canada"},
    "604434": {"jurisdiction": "Northwest Territories", "abbr": "NT", "country": "Canada"},
    "636000": {"jurisdiction": "Virginia", "abbr": "VA", "country": "USA"},
    "636001": {"jurisdiction": "New York", "abbr": "NY", "country": "USA"},
    "636002": {"jurisdiction": "Massachusetts", "abbr": "MA", "country": "USA"},
    "636003": {"jurisdiction": "Maryland", "abbr": "MD", "country": "USA"},
    "636004": {"jurisdiction": "North Carolina", "abbr": "NC", "country": "USA"},
    "636005": {"jurisdiction": "South Carolina", "abbr": "SC", "country": "USA"},
    "636006": {"jurisdiction": "Connecticut", "abbr": "CT", "country": "USA"},
    "636007": {"jurisdiction": "Louisiana", "abbr": "LA", "country": "USA"},
    "636008": {"jurisdiction": "Montana", "abbr": "MT", "country": "USA"},
    "636009": {"jurisdiction": "New Mexico", "abbr": "NM", "country": "USA"},
    "636010": {"jurisdiction": "Florida", "abbr": "FL", "country": "USA"},
    "636011": {"jurisdiction": "Delaware", "abbr": "DE", "country": "USA"},
    "636012": {"jurisdiction": "Ontario", "abbr": "ON", "country": "Canada"},
    "636013": {"jurisdiction": "Nova Scotia", "abbr": "NS", "country": "Canada"},
    "636014": {"jurisdiction": "California", "abbr": "CA", "country": "USA"},
    "636015": {"jurisdiction": "Texas", "abbr": "TX", "country": "USA"},
    "636016": {"jurisdiction": "Newfoundland", "abbr": "NF", "country": "Canada"},
    "636017": {"jurisdiction": "New Brunswick", "abbr": "NB", "country": "Canada"},
    "636018": {"jurisdiction": "Iowa", "abbr": "IA", "country": "USA"},
    "636019": {"jurisdiction": "Guam", "abbr": "GU", "country": "USA"},
    "636020": {"jurisdiction": "Colorado", "abbr": "GM", "country": "USA"},
    "636021": {"jurisdiction": "Arkansas", "abbr": "AR", "country": "USA"},
    "636022": {"jurisdiction": "Kansas", "abbr": "KS", "country": "USA"},
    "636023": {"jurisdiction": "Ohio", "abbr": "OH", "country": "USA"},
    "636024": {"jurisdiction": "Vermont", "abbr": "VT", "country": "USA"},
    "636025": {"jurisdiction": "Pennsylvania", "abbr": "PA", "country": "USA"},
    "636026": {"jurisdiction": "Arizona", "abbr": "AZ", "country": "USA"},
   # "636027": {"jurisdiction": "State Dept. (Diplomatic)", "abbr": "", "country": "USA"},
    "636028": {"jurisdiction": "British Columbia", "abbr": "BC", "country": "Canada"},
    "636029": {"jurisdiction": "Oregon", "abbr": "OR", "country": "USA"},
    "636030": {"jurisdiction": "Missouri", "abbr": "MO", "country": "USA"},
    "636031": {"jurisdiction": "Wisconsin", "abbr": "WI", "country": "USA"},
    "636032": {"jurisdiction": "Michigan", "abbr": "MI", "country": "USA"},
    "636033": {"jurisdiction": "Alabama", "abbr": "AL", "country": "USA"},
    "636034": {"jurisdiction": "North Dakota", "abbr": "ND", "country": "USA"},
    "636035": {"jurisdiction": "Illinois", "abbr": "IL", "country": "USA"},
    "636036": {"jurisdiction": "New Jersey", "abbr": "NJ", "country": "USA"},
    "636037": {"jurisdiction": "Indiana", "abbr": "IN", "country": "USA"},
    "636038": {"jurisdiction": "Minnesota", "abbr": "MN", "country": "USA"},
    "636039": {"jurisdiction": "New Hampshire", "abbr": "NH", "country": "USA"},
    "636040": {"jurisdiction": "Utah", "abbr": "UT", "country": "USA"},
    "636041": {"jurisdiction": "Maine", "abbr": "ME", "country": "USA"},
    "636042": {"jurisdiction": "South Dakota", "abbr": "SD", "country": "USA"},
    "636043": {"jurisdiction": "District of Columbia", "abbr": "DC", "country": "USA"},
    "636044": {"jurisdiction": "Saskatchewan", "abbr": "SK", "country": "Canada"},
    "636045": {"jurisdiction": "Washington", "abbr": "WA", "country": "USA"},
    "636046": {"jurisdiction": "Kentucky", "abbr": "KY", "country": "USA"},
    "636047": {"jurisdiction": "Hawaii", "abbr": "HI", "country": "USA"},
    "636048": {"jurisdiction": "Manitoba", "abbr": "MB", "country": "Canada"},
    "636049": {"jurisdiction": "Nevada", "abbr": "NV", "country": "USA"},
    "636050": {"jurisdiction": "Idaho", "abbr": "ID", "country": "USA"},
    "636051": {"jurisdiction": "Mississippi", "abbr": "MS", "country": "USA"},
    "636052": {"jurisdiction": "Rhode Island", "abbr": "RI", "country": "USA"},
    "636053": {"jurisdiction": "Tennessee", "abbr": "TN", "country": "USA"},
    "636054": {"jurisdiction": "Nebraska", "abbr": "NE", "country": "USA"},
    "636055": {"jurisdiction": "Georgia", "abbr": "GA", "country": "USA"},
}
def get_iin_by_state(abbr):
    for iin, info in IIN_JURISDICTIONS.items():
        if info['abbr'].upper() == abbr.upper():
            return iin
    return None

def ensure_dirs():
    try:
        os.makedirs(BARCODE_DIR, exist_ok=True)
        os.makedirs(DATA_DIR, exist_ok=True)
        os.makedirs(CARDS_DIR, exist_ok=True)
    except Exception as e:
        raise RuntimeError(f"Fatal error: Unable to create necessary directories. {e}")

def format_date(d): return d.strftime("%m%d%Y")

def generate_state_license_number(state):
    """Generate a license number conforming to the state's specific format using faker."""

    state_formats = {
        'AL': lambda: fake.bothify(text='#' * random.randint(1, 7)),
        'AK': lambda: fake.bothify(text='#' * random.randint(1, 7)),
        'AZ': lambda: random.choice([
            fake.bothify(text='?'+('#'*random.randint(1,8))),
            fake.bothify(text='??'+('#'*random.randint(2,5))),
            fake.bothify(text='#'*9)
        ]),
        'AR': lambda: fake.bothify(text='#'*random.randint(4,9)),
        'CA': lambda: fake.bothify(text='?'+'#######'),
        'CO': lambda: random.choice([
            fake.bothify(text='#########'),
            fake.bothify(text='?'+'#'*random.randint(3,6)),
            fake.bothify(text='??'+'#'*random.randint(2,5))
        ]),
        'CT': lambda: fake.bothify(text='#########'),
        'DE': lambda: fake.bothify(text='#'*random.randint(1,7)),
        'DC': lambda: fake.bothify(text='#'*random.choice([7,9])),
        'FL': lambda: fake.bothify(text='?'+'############'),
        'GA': lambda: fake.numerify(text='%######').zfill(9),
        'HI': lambda: random.choice([
            fake.bothify(text='?'+'#'*8),
            fake.bothify(text='#'*9)
        ]),
        'ID': lambda: random.choice([
            fake.bothify(text='??######?'),
            fake.bothify(text='#'*9)
        ]),
        'IL': lambda: random.choice([
            fake.bothify(text='?'+'#'*11),
            fake.bothify(text='?'+'#'*12)
        ]),
        'IN': lambda: random.choice([
            fake.bothify(text='?'+'#'*9),
            fake.bothify(text='#'*9),
            fake.bothify(text='#'*10)
        ]),
        'IA': lambda: random.choice([
            fake.bothify(text='#'*9),
            fake.bothify(text='###??####')
        ]),
        'KS': lambda: random.choice([
            fake.bothify(text='?#?#?'),
            fake.bothify(text='?'+'#'*8),
            fake.bothify(text='#'*9)
        ]),
        'KY': lambda: random.choice([
            fake.bothify(text='?'+'#'*8),
            fake.bothify(text='?'+'#'*9),
            fake.bothify(text='#'*9)
        ]),
        'LA': lambda: fake.bothify(text='#'*random.randint(1,9)),
        'ME': lambda: random.choice([
            fake.bothify(text='#'*7),
            fake.bothify(text='#'*7+'?'),
            fake.bothify(text='#'*8)
        ]),
        'MD': lambda: fake.bothify(text='?'+'#'*12),
        'MA': lambda: random.choice([
            fake.bothify(text='?'+'#'*8),
            fake.bothify(text='#'*9)
        ]),
        'MI': lambda: random.choice([
            fake.bothify(text='?'+'#'*10),
            fake.bothify(text='?'+'#'*12)
        ]),
        'MN': lambda: fake.bothify(text='?'+'#'*12),
        'MS': lambda: fake.bothify(text='#'*9),
        'MO': lambda: random.choice([
            fake.bothify(text='?'+'#'*random.randint(5,9)),
            fake.bothify(text='?'+'#'*6+'R'),
            fake.bothify(text='#'*8+'??'),
            fake.bothify(text='#'*9+'?'),
            fake.bothify(text='#'*9)
        ]),
        # Add more states as needed...
        'NY': lambda: random.choice([
            fake.bothify(text='?'+'#'*7),
            fake.bothify(text='?'+'#'*18),
            fake.bothify(text='#'*8),
            fake.bothify(text='#'*9),
            fake.bothify(text='#'*16),
            fake.lexify(text='????????')
        ]),
        'TX': lambda: fake.bothify(text='#'*random.choice([7,8])),
        'VA': lambda: random.choice([
            fake.bothify(text='?'+'#'*9),
            fake.bothify(text='?'+'#'*10),
            fake.bothify(text='?'+'#'*11),
            fake.bothify(text='#'*9)
        ]),
        'WI': lambda: fake.bothify(text='?'+'#'*13),
        'WY': lambda: fake.bothify(text='#'*random.randint(9,10)),
    }

    return state_formats.get(state, lambda: fake.bothify(text='#'*9))()

def generate_state_subfile(dlid_data: dict, custom_fields: dict) -> dict:
    """Generate a state-specific subfile with custom fields."""
    # this is largely a stub. We want at least 2 subfiles to test parsing
    # dl/id's with more than 1 subfile, but we don't have any data for any
    # specific state. later, we should allow import of JSON specification files
    state_abbr = dlid_data.get("DAJ", "").upper()
    if state_abbr is None:
        raise ValueError("State abbreviation (DAJ) is required in dlid_data.")
    #get first letter of state abbr
    z_label = "Z" + state_abbr[0]  # Use first letter for subfile label
    # if custom_fields is empty, use default test fields
    if not custom_fields:
        county = fake.last_name_male().upper()
        county_label = z_label + "W"  
        test_label = z_label + "T"  
        subfile_data = {
            "subfile_type": z_label,
            county_label: county,
            test_label: "TEST STRING",
            z_label + "X": fake.bothify(text='?' + '#' * random.randint(1, 5)),
        }
    else:
        raise ValueError("Custom fields are not implemented yet.")
        #subfile_data = {k: v for k, v in sorted(custom_fields.items())}

    return subfile_data

def generate_license_data(state=None):
    if state is None:
        state = fake.state_abbr()
    else:
        state = state.upper()
    
    dob = fake.date_of_birth(minimum_age=16, maximum_age=90)
    issue_date = datetime.today()
    exp_date = issue_date + timedelta(days=random.randint(365 * 5, 365 * 10))
    
    # Generate sex first, then use it to determine gender-appropriate names
    sex = random.choice(["1", "2"])  # "1" = male, "2" = female
    
    # Generate gender-appropriate first and middle names
    if sex == "1":  # Male
        first_name = fake.first_name_male().upper()
        middle_name = fake.first_name_male().upper()
    else:  # Female
        first_name = fake.first_name_female().upper()
        middle_name = fake.first_name_female().upper()
    
    eye = random.choice(["BLK", "BLU", "BRO", "GRY", "GRN", "HAZ", "MAR", "PNK", "DIC", "UNK"])
    hair = random.choice(["BLK", "BLN", "BRO", "GRY", "RED", "WHI", "SDY", "UNK"])
    race = random.choice(["W", "B", "A", "I", "U"])
    height = str(random.randint(58, 78))
    weight = str(random.randint(115, 275))
    veteran = random.choice(["0", "1"])
    organ_donor = random.choice(["0", "1"])
    truncation_family_name = random.choice(["N","T","U"]) #not, trunc, unk
    truncation_first_name = random.choice(["N", "T", "U"]) #not, trunc, unk
    truncation_middle_name = random.choice(["N", "T", "U"]) #not, trunc, unk
    dhs_compliance_type = random.choice(["N", "F"]) #F = comp, N = non
    limited_duration_document = random.choice(["1", "0"]) #cardholder temporary lawful status
    country_of_issuance = "USA"
    state_specific_vehicle_class = 'D'
    state_specific_restrictions = ''
    state_specific_endorsements = ''
    
    # Generate state-specific license number
    license_number = generate_state_license_number(state)
    
    dlid_data = {
        "subfile_type": "DL",
        "DAQ": license_number,  # State-specific format
        "DCA": state_specific_vehicle_class,
        "DCB": state_specific_restrictions,
        "DCD": state_specific_endorsements,
        "DBA": format_date(exp_date),
        "DCS": fake.last_name().upper(),
        "DAC": first_name,  # Gender-appropriate first name
        "DAD": middle_name,  # Gender-appropriate middle name
        "DBD": format_date(issue_date),
        "DBB": format_date(dob),
        "DBC": sex,
        "DAY": eye,
        "DAU": height,
        "DAW": weight,
        "DAZ": hair,
        "DCL": race,
        "DAG": fake.street_address().upper(),
        "DAI": fake.city().upper(),
        "DAJ": state,  # Use the specified state
        "DAK": fake.zipcode().replace("-", "").ljust(9, "0"),
        "DCF": fake.unique.bothify(text="DOC#####"),
        "DCG": country_of_issuance,
        "DDE": truncation_family_name,
        "DDF": truncation_first_name,
        "DDG": truncation_middle_name,
        "DDA": dhs_compliance_type,
        "DDB": format_date(issue_date),
        "DDC": format_date(exp_date),
        "DDD": limited_duration_document,
        "DDK": organ_donor,
        "DDL": veteran,
    }
    state_data = generate_state_subfile(dlid_data, custom_fields={})

    subfile_data = [
        dlid_data,
        state_data
    ]

    return subfile_data

def format_barcode_data(data):
    compliance = "@\n\x1E\r"  # @LF RS CR
    file_type = "ANSI "
    version = "10"
    jurisdiction_version = "00"
    # data contains our subfiles. number of entries is number of subfiles
    number_of_entries = f"{len(data):02d}"

    dl_data = data[0]
    state_data = data[1]
    IIN = get_iin_by_state(dl_data["DAJ"]) or "636026"  # Default to Arizona if state not found
    header_base = (
        compliance +
        file_type +
        IIN +
        version +
        jurisdiction_version +
        number_of_entries 
    )
    # Build subfiles 
    # Note, the following should be turned into loop over data
    # we shove DAQ in first, because this is common in most AAMVA barcodes
    daq = dl_data["DAQ"]
    dl_subfile_type = dl_data["subfile_type"]
    dl_fields = {k: v for k, v in dl_data.items() if k != "DAQ" and k != "subfile_type"}
    dl_subfile_data = dl_subfile_type + f"DAQ{daq}\n" +"".join(f"{k}{v}\n" for k, v in dl_fields.items()) + "\r"
    dl_subfile_length = len(dl_subfile_data.encode("ascii"))
    state_fields = {k: v for k, v in state_data.items() if k != "subfile_type"}
    state_subfile_type = state_data["subfile_type"]
    state_subfile_data = state_subfile_type + "".join(f"{k}{v}\n" for k, v in state_fields.items()) + "\r"
    state_subfile_length = len(state_subfile_data.encode("ascii"))

     # Offset needs to be calculated based on the header length and subfile length
    # The offset is the length of the header plus the length of ALL subfile definitions
    # plus all other data before the subfile, which starts at the subfile_type characters (ex: "DL" )
    subfiles_designators_len = 10 * number_of_entries
    dl_subfile_offset = len((header_base + subfiles_designators_len).encode("ascii"))
    dl_subfile_length = len(dl_subfile_data.encode("ascii"))
    dl_subfile_designator = dl_subfile_type + f"{dl_subfile_offset:04d}" + f"{dl_subfile_length:04d}"
    state_subfile_offset = dl_subfile_offset + dl_subfile_length
    state_subfile_designator = state_subfile_type + f"{state_subfile_offset:04d}" + f"{state_subfile_length:04d}"

    header = (
        header_base +
        f"{dl_subfile_designator}" +
        f"{state_subfile_designator}"
    )

    return header + dl_subfile_data + state_subfile_data

def save_barcode_and_data(data, index):
    raw = format_barcode_data(data)
    codes = pdf417.encode(raw, columns=13, security_level=5)
    image = pdf417.render_image(codes)
    bmp_path = os.path.join(BARCODE_DIR, f"license_{index}.bmp")
    txt_path = os.path.join(DATA_DIR, f"license_{index}.txt")
    image.save(bmp_path)
    with open(txt_path, "w") as f: f.write(raw)
    return bmp_path, data

def create_avery_pdf(data_list):
    """Create a PDF using Avery 28371 business card template layout (10 cards per page)"""
    pdf_path = os.path.join(OUTPUT_DIR, "licenses_avery_28371.pdf")
    
    # Avery 28371 specifications
    page_width, page_height = letter  # 8.5" x 11"
    card_width = 3.5 * inch
    card_height = 2 * inch
    
    # Margins and spacing
    left_margin = 0.75 * inch
    top_margin = 0.5 * inch
    horizontal_spacing = 0.25 * inch
    vertical_spacing = 0 * inch
    
    # Create PDF canvas
    c = canvas.Canvas(pdf_path, pagesize=letter)
    
    # Process each card
    for page_num in range(0, len(data_list), 10):
        page_cards = data_list[page_num:page_num + 10]
        
        for card_index, (img_path, data) in enumerate(page_cards):
            # Calculate position on page (2 columns, 5 rows)
            row = card_index // 2
            col = card_index % 2
            
            x = left_margin + col * (card_width + horizontal_spacing)
            y = page_height - top_margin - (row + 1) * (card_height + vertical_spacing)
            
            # Draw card border (optional - remove if not needed)
            c.setStrokeColorRGB(0.8, 0.8, 0.8)
            c.setLineWidth(0.5)
            c.rect(x, y, card_width, card_height, stroke=1, fill=0)
            
            # Add barcode image
            barcode_width = 1.8 * inch
            barcode_height = 0.6 * inch
            barcode_x = x + 0.1 * inch
            barcode_y = y + card_height - barcode_height - 0.1 * inch
            
            # Draw barcode
            try:
                c.drawImage(img_path, barcode_x, barcode_y, 
                           width=barcode_width, height=barcode_height)
            except:
                print(f"Warning: Could not add barcode image {img_path}")
            
            # Add text information
            text_x = x + 0.1 * inch
            text_y = y + card_height - barcode_height - 0.25 * inch
            
            c.setFont("Helvetica", 8)
            c.setFillColorRGB(0, 0, 0)
            dl_data = data[0]  
            state_data = data[1]
            state_line = "|".join([f"{key} {value}" for key, value in state_data.items() if key != "subfile_type"])
            # Format text lines
            lines = [
                f"{dl_data['DAC']} {dl_data['DAD']} {dl_data['DCS']}",
                f"DOB: {dl_data['DBB']} | EXP: {dl_data['DBA']}",
                f"DL#: {dl_data['DAQ']}",
                f"Class: {dl_data['DCA']} | {dl_data['DAI']}, {dl_data['DAJ']}",
                f"{'M' if dl_data['DBC']=='1' else 'F'} | {dl_data['DAY']} | {dl_data['DAZ']} | {dl_data['DAU']}\" | {dl_data['DAW']}lbs",
                f"Organ Donor: {dl_data['DDK']} | Veteran: {dl_data['DDL']}",
                f"{state_line}"
            ]
            
            # Draw text lines
            line_height = 0.15 * inch
            for i, line in enumerate(lines):
                c.drawString(text_x, text_y - i * line_height, line)
        
        # Start new page if there are more cards
        if page_num + 10 < len(data_list):
            c.showPage()
    
    # Save the PDF
    c.save()
    print(f"✅ PDF saved to {pdf_path}")
    return pdf_path

def generate_individual_card_image(data, img_path, width_inches=3.5, dpi=300):
    """Generate an individual card image for use in ODT/DOCX
    
    Args:
        data: array of License data dictionary
        img_path: Path to the barcode image
        width_inches: Desired width in inches (default 3.5" for business cards)
        dpi: Dots per inch for the output image (default 300)
    
    Returns:
        Path to the generated card image
    """
    card_width = int(width_inches * dpi)
    card_height = int(2.0 * dpi)  # 2 inches height for business card
    
    card = PILImage.new("RGB", (card_width, card_height), "white")
    draw = ImageDraw.Draw(card)

    # font size in pixels 
    base_font_size = 50
    small_font_size = 40
    
    # Load font
    try:
        font = ImageFont.truetype("LiberationMono-Bold.ttf", base_font_size)
        small_font = ImageFont.truetype("LiberationMono-Bold.ttf", small_font_size)
    except:
        font = ImageFont.load_default()
        small_font = font
    
    # Add barcode - scale proportionally - mode should be 1 : 
    # 1bit pixes, black and white.
    # or 'L' for grayscale
    barcode_img = PILImage.open(img_path).convert("L")
    barcode_width = int(card_width * 0.55)  
    barcode_height = int(barcode_width * 0.25)  # Maintain barcode aspect ratio
    
    barcode_img = barcode_img.resize((barcode_width, barcode_height), PILImage.Resampling.LANCZOS)
    
    # Position barcode
    barcode_x = int(card_width * 0.03)  # 3% margin
    barcode_y = int(card_height * 0.05)  # 5% margin
    card.paste(barcode_img, (barcode_x, barcode_y))
    
    # Add text
    text_x = barcode_x
    text_y = barcode_y + barcode_height + int(card_height * 0.02)
    line_height = small_font_size
    dl_data = data[0]
    state_data = data[1]
    state_line = "|".join([f"{key} {value}" for key, value in state_data.items() if key != "subfile_type"]) 
    lines = (
            f"{dl_data['DAC']} {dl_data['DAD']} {dl_data['DCS']}\n" + 
            f"DOB: {dl_data['DBB']} | EXP: {dl_data['DBA']}\n" + 
            f"DL#: {dl_data['DAQ']}\n" + 
            f"Class: {dl_data['DCA']} | {dl_data['DAI']}, {dl_data['DAJ']}\n" + 
            f"{'M' if dl_data['DBC']=='1' else 'F'} | {dl_data['DAY']} | {dl_data['DAZ']} | {dl_data['DAU']}\" | {dl_data['DAW']}lbs\n" + 
            f"Organ Donor: {dl_data['DDK']} | Veteran: {dl_data['DDL']}\n" + 
            f"{state_line}"
        )

    draw.text((text_x, text_y), lines, fill="black", font=small_font, spacing=10)    
    # Save the card image
    card_img_path = img_path.replace('.bmp', '_card.png')
    card.save(card_img_path, dpi=(dpi, dpi))
    return card_img_path


def create_odt_card(data_list):
    """
    Create an ODT document with license cards - properly embedding images and text.
    """
    if False:
        from odf.opendocument import OpenDocumentText
        from odf.style import Style, TextProperties
        from odf.text import P
        from odf.draw import Frame, Image as DrawImage

        doc = OpenDocumentText()

        # Define paragraph style for card text
        pstyle = Style(name="CardPara", family="paragraph")
        pstyle.addElement(TextProperties(fontsize="10pt"))
        doc.styles.addElement(pstyle)

        for idx, (img_path, data) in enumerate(data_list):
            dl = data[0]
            # 1) Embed the card image
            frame = Frame(width="3.5in", height="2in", anchortype="as-char")
            img = DrawImage(href=img_path, type="simple", show="embed", actuate="onLoad")
            frame.addElement(img)
            doc.text.addElement(frame)

            # 2) Add human-readable license text beneath
            lines = [
                f"{dl['DAC']} {dl['DAD']} {dl['DCS']}",
                f"DOB: {dl['DBB']} | EXP: {dl['DBA']}",
                f"DL#: {dl['DAQ']}",
                f"Class: {dl['DCA']} | {dl['DAI']}, {dl['DAJ']}",
                f"{'M' if dl['DBC']=='1' else 'F'} | {dl['DAY']} | {dl['DAZ']} | {dl['DAU']}\" | {dl['DAW']}lbs",
                f"Organ Donor: {dl['DDK']} | Veteran: {dl['DDL']}"
            ]
            for line in lines:
                p = P(text=line, stylename=pstyle)
                doc.text.addElement(p)

            # 3) Spacer between cards
            if idx < len(data_list) - 1:
                doc.text.addElement(P())
                doc.text.addElement(P())

        # Save the ODT
        doc.save(ODT_FILE)
        print(f"✅ ODT saved to {ODT_FILE}")
    else:
        print("⚠️ ODT generation disabled. Skipping.")


# === DOCX BUSINESS CARD GENERATION ===
# Requires: pip install python-docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_docx_card(data_list):
    """
    Create a DOCX document using the Avery 28371 business card template layout (2 columns x 5 rows per page).
    Each cell gets a card image that maintains proper aspect ratio.
    """
    doc = Document()
    section = doc.sections[-1]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    card_width = 3.5
    card_height = 2.0
    max_cards_per_page = 10

    def add_card_table():
        table = doc.add_table(rows=5, cols=2)
        table.autofit = False
        
        # Set precise column widths
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                
                # Set cell width
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(card_width * 1440)))  # 1440 twips per inch
                tcW.set(qn('w:type'), 'dxa')
                tc_pr.append(tcW)
                
                # Remove cell margins/padding to maximize space
                tcMar = OxmlElement('w:tcMar')
                for margin_type in ['top', 'left', 'bottom', 'right']:
                    node = OxmlElement(f'w:{margin_type}')
                    node.set(qn('w:w'), '0')
                    node.set(qn('w:type'), 'dxa')
                    tcMar.append(node)
                tc_pr.append(tcMar)
                
                cell.width = Inches(card_width)
            
            # Set row height
            tr_pr = row._tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(int(card_height * 1440)))
            trHeight.set(qn('w:hRule'), 'exact')
            tr_pr.append(trHeight)
            row.height = Inches(card_height)
        
        return table

    table = add_card_table()
    
    for idx, (img_path, data) in enumerate(data_list):
        # Calculate which cell to use
        row_idx = (idx % 10) // 2
        col_idx = idx % 2
        cell = table.cell(row_idx, col_idx)
        
        # Clear any existing content
        cell.text = ""
        paragraph = cell.paragraphs[0]
        
        # Generate high-quality card image
        card_img_path = generate_individual_card_image(data, img_path, width_inches=3.4, dpi=300)
        
        # Add the image with proper sizing
        # We use 3.4" width to leave a small margin in the 3.5" cell
        # Let Word handle the height automatically to maintain aspect ratio
        run = paragraph.add_run()
        run.add_picture(card_img_path, width=Inches(3.4))
        
        # Center the image in the cell
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Remove paragraph spacing
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing = 1.0
        
        # Add new page after every 10 cards (except for the last card)
        if (idx + 1) % max_cards_per_page == 0 and idx != len(data_list) - 1:
            # doc.add_page_break()
            table = add_card_table()
    
    # Save the document
    doc.save(DOCX_FILE)
    print(f"✅ DOCX saved to {DOCX_FILE}")

def main():
    parser = argparse.ArgumentParser(description="Generate AAMVA 2020 DL/ID Licenses")
    parser.add_argument('-a','--all-states', action='store_true',
                        help='Generate one license for each state')
    parser.add_argument('-d', '--directory', type=str, default=OUTPUT_DIR,
                        help='Output directory for generated files (default: output)')
    parser.add_argument('-n', '--number', type=int, default=10)
    parser.add_argument('-s', '--state', type=str, default=None, 
                        help='Two-letter state abbreviation (e.g., CA, FL, TX). If not specified, random states will be used.')
    parser.add_argument('-z', '--state-subfile', action='store_true',
                        help='Generate a state subfile for each card (default: False)')
    parser.add_argument('-p', '--no-pdf', action='store_true',
                        help='Do not generate a PDF for each card (default: False)')
    parser.add_argument('-o', '--no-odt', action='store_true',
                        help='Do not generate an ODT for each card (default: False)')
    parser.add_argument('-w', '--no-word', action='store_true',
                        help='Do not generate a DOCX for each card (default: False)')
    args = parser.parse_args()

    ensure_dirs()
    records = []
    
    if args.all_states:
        # Generate one license for each state
        states = ['AL', 'AK', 'AZ', 'AR', 'CA', 'CO', 'CT', 'DE', 'DC', 'FL', 
                  'GA', 'HI', 'ID', 'IL', 'IN', 'IA', 'KS', 'KY', 'LA', 'ME', 
                  'MD', 'MA', 'MI', 'MN', 'MS', 'MO', 'MT', 'NE', 'NV', 'NH', 
                  'NJ', 'NM', 'NY', 'NC', 'ND', 'OH', 'OK', 'OR', 'PA', 'RI', 
                  'SC', 'SD', 'TN', 'TX', 'UT', 'VT', 'VA', 'WA', 'WV', 'WI', 'WY']
        for i, state in enumerate(states):
            data = generate_license_data(state)
            img_path, d = save_barcode_and_data(data, i)
            records.append((img_path, d))
            print(f"Generated license for {state}: {data[0]['DAQ']}")
    else:
        # Generate specified number of licenses
        for i in range(args.number):
            data = generate_license_data(args.state)
            img_path, d = save_barcode_and_data(data, i)
            records.append((img_path, d))
            print(f"Generated license for {data[0]['DAJ']}: {data[0]['DAQ']}")

    # If no PDF generation is requested, skip it
    if not args.no_pdf:
        print("📄 Generating PDF with Avery 28371 layout...")
        # Create PDF with Avery template layout
        create_avery_pdf(records)     
        print(f"📄 PDF file with Avery 28371 layout: {os.path.join(OUTPUT_DIR, 'licenses_avery_28371.pdf')}")
    
    # If no ODT generation is requested, skip it
    if not args.no_odt:
        print("📄 Generating ODT document with license cards...")
        # Create ODT with individual card images
        create_odt_card(records)
        print(f"📄 ODT file: {ODT_FILE}")

    # If no DOCX generation is requested, skip it
    if not args.no_word:
        print("📄 Generating DOCX document with license cards...")
        # Create DOCX with Avery template layout
        create_docx_card(records)
        print(f"📄 DOCX file: {DOCX_FILE}")

    print("✅ End process.")

if __name__ == "__main__":
    main()
